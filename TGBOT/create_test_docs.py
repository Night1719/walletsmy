#!/usr/bin/env python3
"""
Create test Word documents for Mini App
"""
from docx import Document
import os

def create_test_documents():
    """Create test Word documents"""
    
    # Create directories
    os.makedirs('instructions/1c', exist_ok=True)
    os.makedirs('instructions/email', exist_ok=True)
    
    # 1C AR2 document
    doc = Document()
    doc.add_heading('1C AR2 Инструкция', 0)
    doc.add_paragraph('Это тестовая инструкция для 1C AR2.')
    doc.add_paragraph('Здесь может быть подробное описание работы с системой.')
    doc.add_heading('Шаг 1: Вход в систему', level=1)
    doc.add_paragraph('1. Откройте браузер')
    doc.add_paragraph('2. Перейдите на сайт системы')
    doc.add_paragraph('3. Введите логин и пароль')
    doc.add_heading('Шаг 2: Работа с документами', level=1)
    doc.add_paragraph('• Создание новых документов')
    doc.add_paragraph('• Редактирование существующих')
    doc.add_paragraph('• Отправка на согласование')
    doc.save('instructions/1c/ar2.docx')
    print('✅ Создан: instructions/1c/ar2.docx')
    
    # 1C DM document
    doc2 = Document()
    doc2.add_heading('1C DM Инструкция', 0)
    doc2.add_paragraph('Это тестовая инструкция для 1C DM.')
    doc2.add_paragraph('Инструкция по работе с документооборотом.')
    doc2.add_heading('Основные функции', level=1)
    doc2.add_paragraph('• Создание документов')
    doc2.add_paragraph('• Отправка на согласование')
    doc2.add_paragraph('• Просмотр статуса')
    doc2.add_heading('Настройки системы', level=1)
    doc2.add_paragraph('Для корректной работы системы необходимо:')
    doc2.add_paragraph('1. Настроить права доступа')
    doc2.add_paragraph('2. Создать маршруты согласования')
    doc2.add_paragraph('3. Назначить ответственных')
    doc2.save('instructions/1c/dm.docx')
    print('✅ Создан: instructions/1c/dm.docx')
    
    # iPhone email document
    doc3 = Document()
    doc3.add_heading('Настройка почты iPhone', 0)
    doc3.add_paragraph('Пошаговая инструкция по настройке корпоративной почты на iPhone.')
    doc3.add_heading('Настройка IMAP', level=1)
    doc3.add_paragraph('1. Откройте Настройки')
    doc3.add_paragraph('2. Перейдите в Почта')
    doc3.add_paragraph('3. Добавьте учетную запись')
    doc3.add_heading('Параметры сервера', level=1)
    doc3.add_paragraph('Сервер входящей почты: imap.yourcompany.com')
    doc3.add_paragraph('Порт: 993')
    doc3.add_paragraph('Использовать SSL: Да')
    doc3.save('instructions/email/iphone.docx')
    print('✅ Создан: instructions/email/iphone.docx')
    
    # Android email document
    doc4 = Document()
    doc4.add_heading('Настройка почты Android', 0)
    doc4.add_paragraph('Инструкция по настройке корпоративной почты на Android.')
    doc4.add_heading('Приложение Gmail', level=1)
    doc4.add_paragraph('1. Откройте приложение Gmail')
    doc4.add_paragraph('2. Нажмите "Добавить аккаунт"')
    doc4.add_paragraph('3. Выберите "Другой"')
    doc4.add_heading('Настройки Exchange', level=1)
    doc4.add_paragraph('Сервер: mail.yourcompany.com')
    doc4.add_paragraph('Домен: yourcompany.com')
    doc4.add_paragraph('Использовать SSL: Да')
    doc4.save('instructions/email/android.docx')
    print('✅ Создан: instructions/email/android.docx')
    
    # Outlook document
    doc5 = Document()
    doc5.add_heading('Настройка Outlook', 0)
    doc5.add_paragraph('Инструкция по настройке Microsoft Outlook.')
    doc5.add_heading('Автоматическая настройка', level=1)
    doc5.add_paragraph('1. Запустите Outlook')
    doc5.add_paragraph('2. Введите email адрес')
    doc5.add_paragraph('3. Введите пароль')
    doc5.add_heading('Ручная настройка', level=1)
    doc5.add_paragraph('Если автоматическая настройка не сработала:')
    doc5.add_paragraph('• Сервер входящей почты: outlook.office365.com')
    doc5.add_paragraph('• Порт: 993 (IMAP) или 995 (POP3)')
    doc5.add_paragraph('• Шифрование: SSL/TLS')
    doc5.save('instructions/email/outlook.docx')
    print('✅ Создан: instructions/email/outlook.docx')
    
    print('\n🎉 Все тестовые Word документы созданы!')

if __name__ == "__main__":
    create_test_documents()